import io
import os
from datetime import datetime
from flask import send_file, abort, render_template, current_app
from flask_login import login_required, current_user
from functools import wraps
from auditions import auditions_bp
from auditions.models import Registration, Show, User, AuditionSlot
from auth.decorators import export_required, read_admin_required


def _can_access_show(show_id):
    if not current_user.can_read_admin:
        return False
    if not current_user.managed_shows:
        return True
    return show_id in current_user.managed_shows


# ---------------------------------------------------------------------------
# Excel Export
# ---------------------------------------------------------------------------

@auditions_bp.route('/admin/shows/<int:show_id>/export/xlsx')
@export_required
def export_xlsx(show_id):
    """Export all registrations for a show as an Excel workbook."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        abort(500, 'openpyxl is not installed.')

    show = Show.query.get_or_404(show_id)
    registrations = Registration.query.filter_by(show_id=show.id).filter(
        Registration.status != 'cancelled'
    ).order_by(Registration.status, Registration.created_at).all()

    wb = Workbook()

    # ---- Sheet 1: Registrations ----
    ws = wb.active
    ws.title = 'Registrations'

    header_fill = PatternFill('solid', fgColor='212529')
    header_font = Font(bold=True, color='FFFFFF')
    header_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin = Side(border_style='thin', color='CCCCCC')
    cell_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    # Build headers
    is_admin = current_user.can_evaluate

    base_headers = [
        'Status', 'Last Name', 'First Name', 'Pronouns', 'Email', 'Phone',
        'Audition Date', 'Audition Time',
        'Roles Auditioning For', 'Accept Other Role',
        'Comfortable Performing', 'Equity / ACTRA',
        'Schedule Conflicts', 'Training',
        'Volunteer Interests', 'Video Link',
        'Tags', 'Admin Notes', 'Registered At'
    ]
    if is_admin:
        base_headers.insert(-1, 'Audition Notes')  # before Registered At
    custom_headers = [f['name'] for f in (show.custom_fields or [])]
    all_headers = base_headers + custom_headers

    ws.row_dimensions[1].height = 32
    for col_num, header in enumerate(all_headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = cell_border

    # Status colour map
    status_fills = {
        'confirmed': PatternFill('solid', fgColor='D1FAE5'),
        'waitlisted': PatternFill('solid', fgColor='FEF9C3'),
        'callback':   PatternFill('solid', fgColor='DBEAFE'),
    }

    for row_num, reg in enumerate(registrations, 2):
        u = reg.user
        slot_date = reg.slot.date.strftime('%Y-%m-%d') if reg.slot else ''
        slot_time = reg.slot.start_time.strftime('%I:%M %p') if reg.slot else ''
        tags = ', '.join(t.name for t in reg.tags)
        interests = ', '.join(u.volunteer_interests or [])

        row_data = [
            reg.status.capitalize(),
            u.last_name,
            u.first_name,
            u.pronouns or '',
            u.email,
            u.phone or '',
            slot_date,
            slot_time,
            u.roles_auditioning_for or '',
            'Yes' if u.accept_other_role else 'No',
            'Yes' if u.comfortable_performing else 'No',
            'Yes' if u.equity_or_actra else 'No',
            u.schedule_conflicts or '',
            u.training or '',
            interests,
            reg.video_link or '',
            tags,
            reg.notes or '',
        ]
        if is_admin:
            row_data.append(reg.audition_notes or '')
        row_data.append(reg.created_at.strftime('%Y-%m-%d %H:%M'))

        # Custom field values
        for field in (show.custom_fields or []):
            row_data.append(
                (reg.custom_field_data or {}).get(field['name'], '')
            )

        fill = status_fills.get(reg.status)
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.alignment = Alignment(vertical='top', wrap_text=True)
            cell.border = cell_border
            if fill:
                cell.fill = fill

    # Auto-fit column widths (capped at 50)
    for col_num in range(1, len(all_headers) + 1):
        max_len = 0
        col_letter = get_column_letter(col_num)
        for row in ws.iter_rows(min_col=col_num, max_col=col_num):
            for cell in row:
                try:
                    max_len = max(max_len, len(str(cell.value or '')))
                except Exception:
                    pass
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)

    # ---- Sheet 2: Acting Experience ----
    ws2 = wb.create_sheet('Acting Experience')
    exp_headers = ['Last Name', 'First Name', 'Email', 'Show', 'Role', 'Theatre Group']
    for col_num, header in enumerate(exp_headers, 1):
        cell = ws2.cell(row=1, column=col_num, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = cell_border

    exp_row = 2
    for reg in registrations:
        u = reg.user
        for exp in (u.acting_experience or []):
            row_data = [
                u.last_name, u.first_name, u.email,
                exp.get('show', ''),
                exp.get('role', ''),
                exp.get('theatre_group', '')
            ]
            for col_num, value in enumerate(row_data, 1):
                cell = ws2.cell(row=exp_row, column=col_num, value=value)
                cell.alignment = Alignment(vertical='top')
                cell.border = cell_border
            exp_row += 1

    for col_num in range(1, len(exp_headers) + 1):
        col_letter = get_column_letter(col_num)
        max_len = max(
            (len(str(ws2.cell(r, col_num).value or ''))
             for r in range(1, exp_row)),
            default=10
        )
        ws2.column_dimensions[col_letter].width = min(max_len + 4, 50)

    # ---- Save & send ----
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f"{show.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


# ---------------------------------------------------------------------------
# Word Export
# ---------------------------------------------------------------------------

@auditions_bp.route('/admin/shows/<int:show_id>/export/docx')
@export_required
def export_docx(show_id):
    """Export all registrations for a show as a formatted Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        abort(500, 'python-docx is not installed.')

    show = Show.query.get_or_404(show_id)
    registrations = (
        Registration.query
        .join(Registration.slot)
        .filter(Registration.show_id == show.id)
        .filter(Registration.status != 'cancelled')
        .order_by(AuditionSlot.date, AuditionSlot.start_time, Registration.created_at)
        .all()
    )
    # Also include waitlisted (no slot) at the end
    waitlisted = Registration.query.filter_by(
        show_id=show.id, status='waitlisted'
    ).filter(Registration.slot_id == None).order_by(Registration.created_at).all()  # noqa: E711
    # Remove any waitlisted already caught by the join (those with a slot)
    slotted_ids = {r.id for r in registrations}
    registrations = registrations + [r for r in waitlisted if r.id not in slotted_ids]

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- Title ----
    title = doc.add_heading(show.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x21, 0x25, 0x29)

    sub = doc.add_paragraph(
        f'Audition Registrations  ·  Generated {datetime.now().strftime("%B %d, %Y")}'
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

    # Summary counts
    confirmed = sum(1 for r in registrations if r.status == 'confirmed')
    waitlisted = sum(1 for r in registrations if r.status == 'waitlisted')
    callbacks = sum(1 for r in registrations if r.status == 'callback')

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.add_run(
        f'Confirmed: {confirmed}   ·   Waitlisted: {waitlisted}   ·   Callbacks: {callbacks}'
    ).font.size = Pt(10)

    doc.add_paragraph()

    def _set_cell_bg(cell, hex_color):
        """Set table cell background colour."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _bold_run(para, text, size=10):
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        return run

    def _normal_run(para, text, size=10):
        run = para.add_run(text)
        run.font.size = Pt(size)
        return run

    # Group by status
    groups = [
        ('Callbacks', 'callback', 'D6EAF8'),
        ('Confirmed', 'confirmed', 'D5F5E3'),
        ('Waitlisted', 'waitlisted', 'FEFBD8'),
    ]

    for group_label, status, bg_hex in groups:
        group_regs = [r for r in registrations if r.status == status]
        if not group_regs:
            continue

        # Section heading
        heading = doc.add_heading(f'{group_label} ({len(group_regs)})', level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x21, 0x25, 0x29)
        heading.runs[0].font.size = Pt(13)

        for reg in group_regs:
            u = reg.user

            # Actor name header
            name_para = doc.add_paragraph()
            name_para.paragraph_format.space_before = Pt(8)
            name_para.paragraph_format.space_after = Pt(2)
            _bold_run(name_para, f'{u.first_name} {u.last_name}', size=11)
            if u.pronouns:
                _normal_run(name_para, f'  ({u.pronouns})', size=10)
            if reg.tags:
                tag_str = '  ' + '  '.join(f'[{t.name}]' for t in reg.tags)
                run = name_para.add_run(tag_str)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

            # Details table (2 columns)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            def add_row(label, value):
                if not value:
                    return
                row = table.add_row()
                label_cell = row.cells[0]
                value_cell = row.cells[1]
                _set_cell_bg(label_cell, 'F2F3F4')
                _set_cell_bg(value_cell, 'FFFFFF')
                lp = label_cell.paragraphs[0]
                lp.clear()
                _bold_run(lp, label, size=9)
                vp = value_cell.paragraphs[0]
                vp.clear()
                _normal_run(vp, str(value), size=9)
                # Column widths
                label_cell.width = Cm(4.5)
                value_cell.width = Cm(12)

            # Slot
            if reg.slot:
                slot_str = (f'{reg.slot.date.strftime("%A, %B %d, %Y")}  '
                            f'{reg.slot.start_time.strftime("%I:%M %p")}')
                add_row('Audition Time', slot_str)

            add_row('Email', u.email)
            add_row('Phone', u.phone)
            add_row('Roles Auditioning For', u.roles_auditioning_for)
            add_row('Accept Other Role', 'Yes' if u.accept_other_role else 'No')
            add_row('Comfortable Performing', 'Yes' if u.comfortable_performing else 'No')
            add_row('Equity / ACTRA', 'Yes' if u.equity_or_actra else 'No')
            add_row('Schedule Conflicts', u.schedule_conflicts)
            add_row('Training', u.training)
            add_row('Volunteer Interests', ', '.join(u.volunteer_interests or []))
            add_row('Video Link', reg.video_link)

            # Acting experience
            if u.acting_experience:
                exp_lines = '\n'.join(
                    f'{e.get("show", "")} — {e.get("role", "")} ({e.get("theatre_group", "")})'
                    for e in u.acting_experience
                )
                add_row('Acting Experience', exp_lines)

            # Custom fields
            for field in (show.custom_fields or []):
                val = (reg.custom_field_data or {}).get(field['name'], '')
                add_row(field['name'], val)

            # Admin notes
            if reg.notes:
                add_row('Admin Notes', reg.notes)

            # Audition notes — admin only
            if current_user.can_evaluate and reg.audition_notes:
                add_row('Audition Notes', reg.audition_notes)

            doc.add_paragraph()  # spacer between actors

        doc.add_page_break()

    # Remove trailing page break
    last = doc.paragraphs[-1]
    if last.text == '' and last.runs == []:
        p = last._element
        p.getparent().remove(p)

    # ---- Save & send ----
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{show.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ---------------------------------------------------------------------------
# Audition Sheets — one page per actor (Word + Print/PDF)
# ---------------------------------------------------------------------------

def _sheet_registrations(show):
    """Confirmed + callback registrations in slot date/time order."""
    return (
        Registration.query
        .join(Registration.slot)
        .filter(Registration.show_id == show.id)
        .filter(Registration.status.in_(['confirmed', 'callback']))
        .order_by(AuditionSlot.date, AuditionSlot.start_time, Registration.created_at)
        .all()
    )


@auditions_bp.route('/admin/shows/<int:show_id>/export/sheets/docx')
@export_required
def export_audition_sheets_docx(show_id):
    """One audition sheet per actor as a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        abort(500, 'python-docx is not installed.')

    show = Show.query.get_or_404(show_id)
    if not _can_access_show(show_id):
        abort(403)

    registrations = _sheet_registrations(show)
    if not registrations:
        abort(404, 'No confirmed registrations to export.')

    doc = Document()

    # Letter, narrow margins
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    def _shading(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def _lv(para, label, value, size=9):
        """Add a bold label + normal value run to an existing paragraph."""
        if not value:
            return False
        lr = para.add_run(f'{label}: ')
        lr.bold = True
        lr.font.size = Pt(size)
        vr = para.add_run(str(value))
        vr.font.size = Pt(size)
        return True

    def _row(label, value, size=9):
        """Add a new paragraph with bold label + value."""
        if not value:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        _lv(p, label, value, size)

    for idx, reg in enumerate(registrations):
        u = reg.user

        # ---- Name + pronouns ----
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_before = Pt(0)
        np_.paragraph_format.space_after  = Pt(2)
        nr = np_.add_run(f'{u.first_name} {u.last_name}')
        nr.bold = True
        nr.font.size = Pt(18)
        if u.pronouns:
            pr = np_.add_run(f'  ({u.pronouns})')
            pr.font.size = Pt(11)
            pr.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

        # ---- Show + slot ----
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(3)
        slot_str = ''
        if reg.slot:
            slot_str = (f'{reg.slot.date.strftime("%A, %B %d, %Y")}  ·  '
                        f'{reg.slot.start_time.strftime("%I:%M %p")}')
        sr = sp.add_run(f'{show.title}' + (f'  ·  {slot_str}' if slot_str else ''))
        sr.font.size = Pt(10)
        sr.italic = True
        sr.font.color.rgb = RGBColor(0x21, 0x25, 0x29)

        # ---- Status + tags ----
        stp = doc.add_paragraph()
        stp.paragraph_format.space_after = Pt(6)
        stat_colours = {'confirmed': RGBColor(0x19, 0x87, 0x54),
                        'callback':  RGBColor(0x00, 0x63, 0xcc)}
        str_ = stp.add_run(reg.status.upper())
        str_.bold = True
        str_.font.size = Pt(8)
        if reg.status in stat_colours:
            str_.font.color.rgb = stat_colours[reg.status]
        if reg.tags:
            tr = stp.add_run('   ' + '  '.join(f'[{t.name}]' for t in reg.tags))
            tr.font.size = Pt(8)
            tr.font.color.rgb = RGBColor(0x6c, 0x75, 0x7d)

        # ---- Photo + core details (2-col table if photo exists) ----
        photo_path = None
        if reg.headshot_path:
            candidate = os.path.join(current_app.static_folder, reg.headshot_path)
            if os.path.exists(candidate):
                photo_path = candidate

        if photo_path:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = 'Table Grid'
            lc = tbl.rows[0].cells[0]
            rc = tbl.rows[0].cells[1]
            lc.width = Cm(11)
            rc.width = Cm(5)
            _shading(lc, 'FFFFFF')
            _shading(rc, 'FFFFFF')

            try:
                rp = rc.paragraphs[0]
                rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rp.add_run().add_picture(photo_path, width=Cm(4.5))
            except Exception:
                pass

            def _cell_row(label, value):
                if not value:
                    return
                p = lc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
                lb = p.add_run(f'{label}: ')
                lb.bold = True
                lb.font.size = Pt(9)
                vr = p.add_run(str(value))
                vr.font.size = Pt(9)

            lc.paragraphs[0].clear()
            _cell_row('Email', u.email)
            _cell_row('Phone', u.phone)
            _cell_row('Roles', u.roles_auditioning_for)
            _cell_row('Accept Other Role',    'Yes' if u.accept_other_role else 'No')
            _cell_row('Equity / ACTRA',       'Yes' if u.equity_or_actra else 'No')
            _cell_row('Comfortable Performing','Yes' if u.comfortable_performing else 'No')
            if u.schedule_conflicts:
                _cell_row('Schedule Conflicts', u.schedule_conflicts)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        else:
            _row('Email', u.email)
            _row('Phone', u.phone)
            _row('Roles', u.roles_auditioning_for)
            _row('Accept Other Role',     'Yes' if u.accept_other_role else 'No')
            _row('Equity / ACTRA',        'Yes' if u.equity_or_actra else 'No')
            _row('Comfortable Performing','Yes' if u.comfortable_performing else 'No')
            _row('Schedule Conflicts', u.schedule_conflicts)

        _row('Training', u.training)

        if u.volunteer_interests:
            _row('Volunteer Interests', ', '.join(u.volunteer_interests))

        # ---- Acting experience ----
        if u.acting_experience:
            ep = doc.add_paragraph()
            ep.paragraph_format.space_before = Pt(6)
            ep.paragraph_format.space_after  = Pt(2)
            er = ep.add_run('Acting Experience')
            er.bold = True
            er.font.size = Pt(9)

            exp_tbl = doc.add_table(rows=1, cols=3)
            exp_tbl.style = 'Table Grid'
            for i, h in enumerate(['Show', 'Role', 'Theatre Group']):
                c = exp_tbl.rows[0].cells[i]
                _shading(c, 'F2F3F4')
                cr = c.paragraphs[0].add_run(h)
                cr.bold = True
                cr.font.size = Pt(8)
            for exp in (u.acting_experience or [])[:6]:
                row = exp_tbl.add_row()
                for i, key in enumerate(['show', 'role', 'theatre_group']):
                    c = row.cells[i]
                    c.paragraphs[0].add_run(exp.get(key, '') or '').font.size = Pt(8)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # ---- Custom fields ----
        for field in (show.custom_fields or []):
            val = (reg.custom_field_data or {}).get(field['name'], '')
            _row(field['name'], val)

        # ---- Video link ----
        _row('Video Link', reg.video_link)

        # ---- Admin notes ----
        if reg.notes:
            _row('Notes', reg.notes)

        # ---- Audition notes / blank lines ----
        if current_user.can_evaluate:
            anp = doc.add_paragraph()
            anp.paragraph_format.space_before = Pt(10)
            anp.paragraph_format.space_after  = Pt(3)
            anr = anp.add_run('AUDITION NOTES')
            anr.bold = True
            anr.font.size = Pt(9)

            if reg.audition_notes:
                ap2 = doc.add_paragraph(reg.audition_notes)
                ap2.runs[0].font.size = Pt(9)
            else:
                for _ in range(4):
                    lp = doc.add_paragraph('_' * 90)
                    lp.paragraph_format.space_before = Pt(3)
                    lp.paragraph_format.space_after  = Pt(3)
                    lp.runs[0].font.size = Pt(8)
                    lp.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        if idx < len(registrations) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = (f"{show.title.replace(' ', '_')}_Audition_Sheets_"
                f"{datetime.now().strftime('%Y%m%d')}.docx")
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@auditions_bp.route('/admin/shows/<int:show_id>/export/sheets/print')
@export_required
def export_audition_sheets_print(show_id):
    """Print-optimised HTML audition sheets — use browser Print → Save as PDF."""
    show = Show.query.get_or_404(show_id)
    if not _can_access_show(show_id):
        abort(403)
    registrations = _sheet_registrations(show)
    return render_template(
        'auditions/admin/audition_sheets_print.html',
        show=show,
        registrations=registrations,
        can_evaluate=current_user.can_evaluate,
    )
