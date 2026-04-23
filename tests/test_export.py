"""Tests for Excel and Word export routes."""
import pytest
from tests.conftest import login_admin


@pytest.fixture
def show_with_data(db, admin, actor, slot_show, slot):
    """Show with one confirmed registration, one waitlisted."""
    from auditions.models import Registration, User, Tag

    # Confirmed reg
    r1 = Registration(
        user_id=actor.id,
        show_id=slot_show.id,
        slot_id=slot.id,
        status='confirmed',
        video_link='https://youtube.com/test',
        notes='Great stage presence.',
    )
    slot.current_count = 1

    # Second actor — waitlisted
    u2 = User(
        email='actor2@test.com', first_name='Bob', last_name='Jones',
        role='actor', contact_email_ok=True,
        accept_other_role=True, comfortable_performing=True, equity_or_actra=False,
        acting_experience=[{'show': 'Grease', 'role': 'Danny', 'theatre_group': 'Aurora Drama'}],
        volunteer_interests=['Stagehand', 'Props Master'],
    )
    u2.set_password('pass')
    db.session.add_all([r1, u2])
    db.session.commit()

    r2 = Registration(user_id=u2.id, show_id=slot_show.id, status='waitlisted')
    tag = Tag(name='strong singer')
    db.session.add_all([r2, tag])
    db.session.commit()

    r1.tags.append(tag)
    db.session.commit()

    return slot_show


class TestXlsxExport:

    def test_xlsx_returns_200(self, client, admin, show_with_data):
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx')
        assert r.status_code == 200

    def test_xlsx_content_type(self, client, admin, show_with_data):
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx')
        assert 'spreadsheetml' in r.content_type

    def test_xlsx_is_valid_workbook(self, client, admin, show_with_data):
        import io
        from openpyxl import load_workbook
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx')
        wb = load_workbook(io.BytesIO(r.data))
        assert 'Registrations' in wb.sheetnames
        assert 'Acting Experience' in wb.sheetnames

    def test_xlsx_contains_actor_name(self, client, admin, show_with_data):
        import io
        from openpyxl import load_workbook
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx')
        wb = load_workbook(io.BytesIO(r.data))
        ws = wb['Registrations']
        all_values = [str(cell.value or '') for row in ws.iter_rows() for cell in row]
        assert 'Jane' in all_values

    def test_xlsx_excludes_cancelled(self, client, admin, show_with_data, db):
        import io
        from openpyxl import load_workbook
        from auditions.models import Registration
        # Cancel the confirmed registration
        reg = Registration.query.filter_by(show_id=show_with_data.id, status='confirmed').first()
        reg.status = 'cancelled'
        db.session.commit()
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx')
        wb = load_workbook(io.BytesIO(r.data))
        ws = wb['Registrations']
        all_values = [str(cell.value or '') for row in ws.iter_rows() for cell in row]
        # Jane was cancelled so shouldn't appear
        assert 'Jane' not in all_values

    def test_xlsx_requires_admin(self, client, actor, show_with_data):
        from tests.conftest import login_actor
        login_actor(client, actor)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/xlsx',
                       follow_redirects=False)
        assert r.status_code in (302, 403)


class TestDocxExport:

    def test_docx_returns_200(self, client, admin, show_with_data):
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/docx')
        assert r.status_code == 200

    def test_docx_content_type(self, client, admin, show_with_data):
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/docx')
        assert 'wordprocessingml' in r.content_type

    def test_docx_is_valid_document(self, client, admin, show_with_data):
        import io
        from docx import Document
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/docx')
        doc = Document(io.BytesIO(r.data))
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert show_with_data.title in full_text

    def test_docx_contains_actor_name(self, client, admin, show_with_data):
        import io
        from docx import Document
        login_admin(client, admin)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/docx')
        doc = Document(io.BytesIO(r.data))
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'Jane' in full_text

    def test_docx_requires_admin(self, client, actor, show_with_data):
        from tests.conftest import login_actor
        login_actor(client, actor)
        r = client.get(f'/auditions/admin/shows/{show_with_data.id}/export/docx',
                       follow_redirects=False)
        assert r.status_code in (302, 403)
